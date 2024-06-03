from flask import Flask, request, jsonify
import os
from pathlib import Path
from JsonFileDocument import JsonFileDocument
import PyPDF2
from XmlHtmlDocument import XmlHtmlDocument
from TokenProcessor import TokenProcessor,spanishToken,frenchToken
from indexing import PositionalInvertedIndex,PositionalInvertedIndexSqlite,DiskIndexWriter,DiskPositionalIndex
from docx import Document 
from langdetect import detect
from querying import BooleanQueryParser
import sys
from textblob import TextBlob
from flask_cors import CORS
from porter2stemmer import Porter2Stemmer
import re


app = Flask(__name__)
CORS(app, origins=["http://localhost:3000","http://localhost:3000"])
documents = {}
directories = ["Data/json", "Data/txt", "Data/pdf", "Data/xml", "Data/docx"]

filesJsonDiskSqlite = PositionalInvertedIndexSqlite()





boolean_query_parser=None

def load_pdf(pdf_file_path):
    with open(pdf_file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)  # Use len(reader.pages) to get the number of pages
        body = ""

        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            body += page.extract_text()

        title = os.path.splitext(os.path.basename(pdf_file_path))[0] # Use the file name as the title
        return {"title":title, "body":body, 'filetype':'pdf'}

def LoadDocuments(directory_path):
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                title = os.path.splitext(filename)[0]  # Use the file name as the title
                content = txt_file.read()
                documents.append({'title': title, 'body': content,'filetype':'txt'})
        elif filename.endswith('.json'):
            json_document = JsonFileDocument(file_path)
            documents.append({'title': file_path.split('/')[-1]+" "+json_document.get_title(), 'body': json_document.get_body(), 'url': json_document.get_url(),'filetype':'json'})
        elif filename.endswith(('.xml', '.html', '.htm')):
            xml_html_document = XmlHtmlDocument(file_path)
            documents.append(xml_html_document.dataSend())
        elif filename.endswith('.pdf'):
            pdf_document = load_pdf(file_path)
            documents.append(pdf_document)
        elif filename.endswith('.docx'):
            doc = Document(file_path)
            text_content = " ".join([paragraph.text for paragraph in doc.paragraphs])
            documents.append({'title': os.path.splitext(os.path.basename(file_path))[0], 'body': text_content, 'filetype':'docx'})
    return documents

def GetTokenData(rowdata):
    # Check if rowdata['body'] is empty or None
    if not rowdata['body']:
        return {"fileName": rowdata["title"], "tokenData": [],"indexAndToken":[],"sqlitesDatatoken":[]}
    resulting_types=[]
    detected_language = detect(rowdata['body'])
    print(f"\n\nThe detected language is: {detected_language}\n\n") 

    if True:
        token_processor = TokenProcessor()
        resulting_types = token_processor.process_token(rowdata['body'])
        resulting_terms = [token_processor.normalize_type(t) for t in resulting_types]
        # filtered_terms = token_processor.remove_stopwords(resulting_types)
        return {"fileName": rowdata["title"], "tokenData": resulting_terms , "indexAndToken":token_processor.get_tokens_and_positions(rowdata['body']),"sqlitesDatatoken":token_processor.get_tokens_and_sqlite(rowdata['body'])}
    
    # elif detected_language == "es":
    #     processor = spanishToken()
    #     spanish_tokens, spanish_terms = processor.process_text(rowdata['body'])
    #     return {"fileName": rowdata["title"], "tokenData": spanish_tokens,"indexAndToken":[],"sqlitesDatatoken":[]}
    # elif detected_language == "fr":
    #     processor = frenchToken()
    #     fr_tokens, fr_terms = processor.process_text(rowdata['body'])
    #     return {"fileName": rowdata["title"], "tokenData": fr_tokens,"indexAndToken":[],"sqlitesDatatoken":[]}
    # else:
    #     return {"fileName": rowdata["title"], "tokenData": [],"indexAndToken":[],"sqlitesDatatoken":[]}





def tokenize_with_positions(document):
    if document == []:
        return {}
    positions = {}
    for idx, token in enumerate(document):
        if token in positions:
            positions[token].append(idx)
        else:
            positions[token] = [idx]
    return positions


def load_filesDB():
    for i in range(len(['Data/json'])):     
        documents['Data/json']=LoadDocuments('Data/json')
    for document_id, document in enumerate(documents["Data/json"]):
        data=GetTokenData(document)
        positions = tokenize_with_positions(data["sqlitesDatatoken"])
        for term, term_positions in positions.items():
            filesJsonDiskSqlite.add_term(term, int(data['fileName'].split('.')[0]), term_positions)
    writer = DiskIndexWriter(filesJsonDiskSqlite, "vocab_term_mapping.db", "postings.bin")
    writer.write_index()

    


# with app.app_context():
#         load_filesDB()
   
def convert_text_to_query_format(text):
        result2 = re.findall(r'\w+|AND|OR|NOT|"[^"]+"|[+-]', text)
        stemmer = Porter2Stemmer()
        terms = []
        operations=[]
        
        result=[]
        for i in range(len(result2)):
            if result2[i] == "+":
                result.append("OR")
            elif result2[i] == "-":
                result.append("AND NOT")
            else:
                result.append(stemmer.stem(result2[i]))
        print(result)
        for i in range(len(result)):
            if result[i]== "OR" or result[i]== "AND NOT":
                operations.append(result[i])
            else:
                terms.append(result[i])
                if i+1 <len(result):
                    if result[i+1]!= "OR" and result[i+1]!= "AND NOT":
                        operations.append("AND")
                        
            
        print(terms,operations)
        return terms, operations
    
def convert_text_to_query_formatfor_rankquery(text):
        result2 = re.findall(r'\w+|AND|OR|NOT|"[^"]+"|[+-]', text)
        stemmer = Porter2Stemmer()
        result=[]
        for i in range(len(result2)):
            result.append(stemmer.stem(result2[i]))
        return result

@app.route('/searchdata', methods=['POST'])
def readDisk_string():
    try:
        # Get the JSON data from the request
        data = request.get_json()
        if 'text' not in data:
            return jsonify({'error': 'Missing "text" field in JSON data'}), 400
        text=data['text']
        disk_index = DiskPositionalIndex("vocab_term_mapping.db", "postings.bin")
        terms, operations = convert_text_to_query_format(text)
        result = disk_index.query(terms, operations)
        ids1 = [{"doc_id":str(item[0])+'.json'} for item in result]
        
        if not ids1:
            response_data = {'data':[],"message":"Term not found in any documents","file":0}
        else:
            response_data = {'data':ids1,"message":"done","file":len(ids1)}
        return jsonify(response_data), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
    
    
@app.route('/rankquery', methods=['POST'])
def rankquery():
    try:
        # Get the JSON data from the request
        data = request.get_json()
        if 'text' not in data:
            return jsonify({'error': 'Missing "text" field in JSON data'}), 400
        text=data['text']
        type=data['type']
        
        disk_index = DiskPositionalIndex("vocab_term_mapping.db", "postings.bin")
        terms = convert_text_to_query_formatfor_rankquery(text)
        
        result = disk_index.queryRank(terms,type)
        
        if not result:
            response_data = {'data':[],"message":"Term not found in any documents","file":0}
        else:
            response_data = {'data':result,"message":"done","file":len(result)}
        return jsonify(response_data), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500





